import os
import json
import PyPDF2
import requests
import docx
from pathlib import Path
import time
from tqdm import tqdm

class RegulationApplicabilityAnalyzer:
    def __init__(self, api_key="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"):
        """
        Initialize the analyzer with API key for OpenRouter.
        
        Args:
            api_key (str): OpenRouter API key. Default is the provided key.
        """
        self.api_key = api_key
        
        self.base_dir = Path(r"C:\Users\91810\OneDrive\Desktop\CPC_AIB_Life")
        self.regulations_path = self.base_dir / "output" / "Data_extraction" / "regulation_17A_48_combined.json"
        self.output_dir = self.base_dir / "output" / "regulations_applicable"
        self.client_report_dir = self.base_dir / "Input" / "Client Annual Report"
        
        # Ensure output directory exists
        os.makedirs(self.output_dir, exist_ok=True)
    
    def extract_text_from_pdf(self, pdf_path):
        """Extract text from a PDF file."""
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def find_client_report(self):
        """Find the client annual report in the input directory."""
        for file in os.listdir(self.client_report_dir):
            if file.lower().endswith('.pdf'):
                return self.client_report_dir / file
        raise FileNotFoundError(f"No PDF file found in {self.client_report_dir}")
    
    def load_regulations(self):
        """Load the combined regulations JSON file."""
        with open(self.regulations_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    def analyze_regulation(self, regulation, client_report_text):
        """
        Analyze a single regulation against the client report using Claude 3.7 Sonnet.
        
        Args:
            regulation (dict): The regulation to analyze
            client_report_text (str): The text of the client annual report
            
        Returns:
            dict: The regulation with added applicability fields
        """
        # Prepare the prompt for Claude
        prompt = f"""
You are an expert financial regulations analyst. Your task is to determine if a specific regulation applies to a company based on their annual report.

# Regulation Details:
- Source Name: {regulation['Source Name']}
- Section Type: {regulation['Section Type']}
- Part Number: {regulation['Part Number']}
- Part Name: {regulation['Part Name']}
- Chapter Number: {regulation['Chapter Number']}
- Chapter Name: {regulation['Chapter Name']}
- Regulation Number: {regulation['Regulation Number']}
- Regulation Title: {regulation['Regulation Title']}
- Regulation Text: {regulation['Regulation Text']}

# Context on the Regulations:
The CPC Regulations 2026 outline a comprehensive framework for the conduct of financial service providers in Ireland, combining two key regulatory instruments: the Central Bank Reform Act 2010 (Section 17A) and the Supervision and Enforcement Act 2013 (Section 48). Together, these establish the Standards for Business and the Consumer Protection Regulations applicable from March 24, 2026.

Section 17A defines high-level standards for regulated entities in their dealings with customers. It requires firms to act honestly, fairly, and professionally, securing customers' interests, providing clear and effective information, and ensuring robust governance, risk management, and financial resilience. It includes specific obligations related to managing financial abuse, maintaining adequate financial resources, disclosing material information to regulators, and ensuring systems and controls (including outsourced functions) support ongoing compliance.

Section 48 sets out more detailed consumer protection rules. It applies to entities conducting regulated activities with consumers, including personal and vulnerable consumers. It covers transparency in advertising, fair analysis of the market, handling complaints, providing appropriate advice, avoiding conflicts of interest, and safeguarding consumer assets. Extensive definitions are provided to clarify terms like "consumer," "financial services," and "regulated entity," aligning with EU directives such as Solvency II and MiFID.

Certain exemptions apply. For instance, the regulations may not apply to credit unions (except as insurance intermediaries), crowdfunding platforms regulated under EU law, or reinsurance businesses. Regulations also distinguish between regulated and unregulated activities, emphasizing the importance of clarity in customer communication.

# Client Annual Report:
```
{client_report_text[:10000]}  # Limiting to first 10,000 chars to avoid token limits
```

Based on the annual report and the regulation details, determine:

1. Applicability: Does this regulation apply to this company? Answer with EXACTLY ONE of these options:
   - "Applies"
   - "Does Not Apply"
   - "May Apply - Requires Further Review"

2. Reason for Applicability: Provide a concise reason (maximum 30 words) explaining why this regulation applies or doesn't apply to this company.

3. How are you reasoning: Specify which section(s) of the annual report you analyzed to make this determination.

4. Confidence Score: Rate your confidence in this assessment on a scale of:
   - "Very High" (90-100% confident)
   - "High" (70-89% confident)
   - "Medium" (50-69% confident)
   - "Low" (30-49% confident)
   - "Very Low" (0-29% confident)

Format your response as a JSON object with these exact keys:
{{
  "Applicability": "your answer",
  "Reason for Applicability": "your reason",
  "How are you reasoning": "sections analyzed",
  "Confidence Score": "your confidence level"
}}
"""

        # Make API call to OpenRouter with Claude 3.7 Sonnet
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "anthropic/claude-3-7-sonnet",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.1,  # Low temperature for more deterministic responses
            "max_tokens": 1000
        }
        
        try:
            response = requests.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers=headers,
                json=data
            )
            response.raise_for_status()
            
            # Extract the response content
            result = response.json()
            response_text = result["choices"][0]["message"]["content"]
            
            # Extract the JSON part from the response
            try:
                # Find JSON in the response
                start_idx = response_text.find('{')
                end_idx = response_text.rfind('}') + 1
                if start_idx >= 0 and end_idx > start_idx:
                    json_str = response_text[start_idx:end_idx]
                    analysis_result = json.loads(json_str)
                else:
                    # Fallback if JSON not found
                    analysis_result = {
                        "Applicability": "May Apply - Requires Further Review",
                        "Reason for Applicability": "Could not parse response from AI",
                        "How are you reasoning": "Error in response format",
                        "Confidence Score": "Low"
                    }
            except json.JSONDecodeError:
                # Fallback if JSON parsing fails
                analysis_result = {
                    "Applicability": "May Apply - Requires Further Review",
                    "Reason for Applicability": "Could not parse response from AI",
                    "How are you reasoning": "Error in response format",
                    "Confidence Score": "Low"
                }
            
            # Update the regulation with the analysis results
            regulation.update(analysis_result)
            return regulation
            
        except Exception as e:
            print(f"Error analyzing regulation {regulation['Regulation Number']}: {str(e)}")
            # Return regulation with error information
            regulation.update({
                "Applicability": "May Apply - Requires Further Review",
                "Reason for Applicability": f"Error during analysis: {str(e)[:30]}",
                "How are you reasoning": "Error occurred during API call",
                "Confidence Score": "Low"
            })
            return regulation
    
    def analyze_all_regulations(self, client_report_path=None):
        """
        Analyze all regulations against the client report.
        
        Args:
            client_report_path (Path, optional): Path to the client report. If None, will try to find it.
            
        Returns:
            list: List of regulations with applicability analysis
        """
        # Find client report if not provided
        if client_report_path is None:
            client_report_path = self.find_client_report()
        
        # Extract text from client report
        client_report_text = self.extract_text_from_pdf(client_report_path)
        
        # Load regulations
        regulations = self.load_regulations()
        
        # Analyze each regulation
        analyzed_regulations = []
        for regulation in tqdm(regulations, desc="Analyzing regulations"):
            # Add a small delay to avoid rate limiting
            time.sleep(1)
            analyzed_regulation = self.analyze_regulation(regulation, client_report_text)
            analyzed_regulations.append(analyzed_regulation)
        
        return analyzed_regulations
    
    def save_analyzed_regulations(self, analyzed_regulations):
        """
        Save the analyzed regulations to a JSON file.
        
        Args:
            analyzed_regulations (list): List of regulations with applicability analysis
            
        Returns:
            Path: Path to the saved JSON file
        """
        output_path = self.output_dir / "analyzed_regulations.json"
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(analyzed_regulations, f, indent=2, ensure_ascii=False)
        
        print(f"Analyzed regulations saved to {output_path}")
        return output_path
    
    def generate_summary_document(self, analyzed_regulations):
        """
        Generate a summary Word document of applicable/non-applicable regulations.
        
        Args:
            analyzed_regulations (list): List of regulations with applicability analysis
            
        Returns:
            Path: Path to the saved Word document
        """
        # Create a new Word document
        doc = docx.Document()
        
        # Add title
        doc.add_heading('Regulation Applicability Analysis Summary', 0)
        
        # Add introduction
        doc.add_paragraph('This document summarizes the applicability of Central Bank regulations to the company based on analysis of the annual report.')
        
        # Group regulations by applicability
        applies = [r for r in analyzed_regulations if r.get('Applicability') == 'Applies']
        does_not_apply = [r for r in analyzed_regulations if r.get('Applicability') == 'Does Not Apply']
        may_apply = [r for r in analyzed_regulations if r.get('Applicability') == 'May Apply - Requires Further Review']
        
        # Function to add regulations to document
        def add_regulations_section(title, regulations):
            doc.add_heading(title, 1)
            
            if not regulations:
                doc.add_paragraph('No regulations in this category.')
                return
            
            # Group by Section Type and Part Number
            by_section = {}
            for reg in regulations:
                key = f"{reg['Section Type']} - Part {reg['Part Number']}: {reg['Part Name']}"
                if key not in by_section:
                    by_section[key] = []
                by_section[key].append(reg)
            
            # Add each section
            for section_key, section_regs in by_section.items():
                doc.add_heading(section_key, 2)
                
                for reg in section_regs:
                    # Add regulation details
                    p = doc.add_paragraph()
                    p.add_run(f"Regulation {reg['Regulation Number']}: {reg['Regulation Title']}").bold = True
                    doc.add_paragraph(f"Reason: {reg.get('Reason for Applicability', 'Not specified')}")
                    doc.add_paragraph(f"Analysis based on: {reg.get('How are you reasoning', 'Not specified')}")
                    doc.add_paragraph(f"Confidence: {reg.get('Confidence Score', 'Not specified')}")
                    
                    # Add a separator
                    doc.add_paragraph('-' * 40)
        
        # Add sections for each applicability category
        add_regulations_section('Regulations That Apply', applies)
        add_regulations_section('Regulations That May Apply - Require Further Review', may_apply)
        add_regulations_section('Regulations That Do Not Apply', does_not_apply)
        
        # Add statistics
        doc.add_heading('Summary Statistics', 1)
        doc.add_paragraph(f"Total regulations analyzed: {len(analyzed_regulations)}")
        doc.add_paragraph(f"Regulations that apply: {len(applies)} ({len(applies)/len(analyzed_regulations)*100:.1f}%)")
        doc.add_paragraph(f"Regulations that may apply: {len(may_apply)} ({len(may_apply)/len(analyzed_regulations)*100:.1f}%)")
        doc.add_paragraph(f"Regulations that do not apply: {len(does_not_apply)} ({len(does_not_apply)/len(analyzed_regulations)*100:.1f}%)")
        
        # Save the document
        output_path = self.output_dir / "regulation_applicability_summary.docx"
        doc.save(output_path)
        
        print(f"Summary document saved to {output_path}")
        return output_path
    
    def run_analysis(self):
        """Run the complete analysis pipeline."""
        # Find client report
        client_report_path = self.find_client_report()
        print(f"Found client report: {client_report_path}")
        
        # Analyze regulations
        analyzed_regulations = self.analyze_all_regulations(client_report_path)
        
        # Save analyzed regulations
        json_path = self.save_analyzed_regulations(analyzed_regulations)
        
        # Generate summary document
        doc_path = self.generate_summary_document(analyzed_regulations)
        
        return {
            "json_path": json_path,
            "doc_path": doc_path,
            "regulations_count": len(analyzed_regulations)
        }

if __name__ == "__main__":
    # Run analysis
    analyzer = RegulationApplicabilityAnalyzer()
    results = analyzer.run_analysis()
    
    print("\nAnalysis complete!")
    print(f"Analyzed {results['regulations_count']} regulations")
    print(f"JSON output: {results['json_path']}")
    print(f"Summary document: {results['doc_path']}")
    print("\nNote: Using the provided OpenRouter API key for Claude 3.7 Sonnet")
