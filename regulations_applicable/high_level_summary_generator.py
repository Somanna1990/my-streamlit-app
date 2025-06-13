import os
import json
import time
import requests
import hashlib
from pathlib import Path
from datetime import datetime
import docx
from docx.shared import RGBColor

class HighLevelSummaryGenerator:
    def __init__(self, api_key="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"):
        """
        Initialize the summary generator with API key for OpenRouter.
        
        Args:
            api_key (str): OpenRouter API key. Default is the provided key.
        """
        self.api_key = api_key
        
        self.base_dir = Path(r"C:\Users\91810\OneDrive\Desktop\CPC_AIB_Life")
        self.analyzed_regulations_path = self.base_dir / "output" / "regulations_applicable" / "analyzed_regulations.json"
        self.output_dir = self.base_dir / "output" / "regulations_applicable"
        self.cache_dir = self.base_dir / "output" / "regulations_applicable" / "cache"
        
        # Ensure output and cache directories exist
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.cache_dir, exist_ok=True)
    
    def load_analyzed_regulations(self):
        """Load the analyzed regulations JSON file."""
        with open(self.analyzed_regulations_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    def get_cache_key(self, prompt):
        """Generate a unique cache key for a prompt."""
        return hashlib.md5(prompt.encode('utf-8')).hexdigest()
    
    def check_cache(self, cache_key):
        """Check if a response is cached."""
        cache_file = self.cache_dir / f"{cache_key}.json"
        if cache_file.exists():
            with open(cache_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return None
    
    def save_to_cache(self, cache_key, response):
        """Save a response to the cache."""
        cache_file = self.cache_dir / f"{cache_key}.json"
        with open(cache_file, 'w', encoding='utf-8') as f:
            json.dump(response, f, indent=2, ensure_ascii=False)
    
    def generate_thematic_analysis(self, analyzed_regulations):
        """
        Generate a thematic analysis of the regulations using Claude 3.7 Sonnet.
        
        Args:
            analyzed_regulations (list): The analyzed regulations
            
        Returns:
            dict: The thematic analysis
        """
        # Group regulations by applicability
        applies = [r for r in analyzed_regulations if r.get('Applicability') == 'Applies']
        does_not_apply = [r for r in analyzed_regulations if r.get('Applicability') == 'Does Not Apply']
        may_apply = [r for r in analyzed_regulations if r.get('Applicability') == 'May Apply - Requires Further Review']
        
        # Prepare the prompt
        prompt = f"""
        You are a financial regulatory expert. Based on the following analysis of regulations that apply to a financial institution, 
        provide a high-level thematic analysis of the regulatory landscape for this institution.

        Group the applicable regulations into 5-7 key thematic areas (e.g., consumer protection, governance, reporting requirements, etc.).
        For each thematic area:
        1. Provide a concise summary of the key requirements and obligations
        2. Identify the most significant regulatory risks
        3. Suggest high-level implementation priorities

        Also provide a brief executive summary of the overall regulatory position.

        Here are the statistics:
        - Total regulations analyzed: {len(analyzed_regulations)}
        - Regulations that apply: {len(applies)} ({len(applies)/len(analyzed_regulations)*100:.1f}%)
        - Regulations that may apply: {len(may_apply)} ({len(may_apply)/len(analyzed_regulations)*100:.1f}%)
        - Regulations that do not apply: {len(does_not_apply)} ({len(does_not_apply)/len(analyzed_regulations)*100:.1f}%)

        Here are some examples of regulations that apply:
        {json.dumps([{'Regulation Number': r['Regulation Number'], 
                     'Regulation Title': r['Regulation Title'], 
                     'Reason for Applicability': r.get('Reason for Applicability', 'Not specified')} 
                    for r in applies[:10]], indent=2)}
        
        Your response should be structured, concise, and focused on strategic insights rather than detailed regulation-by-regulation analysis.
        """
        
        # Check cache first
        cache_key = self.get_cache_key(prompt)
        cached_response = self.check_cache(cache_key)
        if cached_response:
            print("Using cached thematic analysis")
            return cached_response
        
        # Make API call to Claude 3.7 Sonnet
        print("Generating thematic analysis using Claude 3.7 Sonnet...")
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "anthropic/claude-3-sonnet-20240229",
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 4000
        }
        
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=data
        )
        
        if response.status_code != 200:
            raise Exception(f"API request failed with status code {response.status_code}: {response.text}")
        
        result = response.json()
        thematic_analysis = result['choices'][0]['message']['content']
        
        # Save to cache
        self.save_to_cache(cache_key, thematic_analysis)
        
        return thematic_analysis
    
    def generate_risk_assessment(self, analyzed_regulations):
        """
        Generate a risk assessment based on the regulations using Claude 3.7 Sonnet.
        
        Args:
            analyzed_regulations (list): The analyzed regulations
            
        Returns:
            dict: The risk assessment
        """
        # Group regulations by applicability
        applies = [r for r in analyzed_regulations if r.get('Applicability') == 'Applies']
        may_apply = [r for r in analyzed_regulations if r.get('Applicability') == 'May Apply - Requires Further Review']
        
        # Prepare the prompt
        prompt = f"""
        You are a financial regulatory risk expert. Based on the following analysis of regulations that apply to a financial institution, 
        provide a high-level risk assessment of the regulatory compliance posture.

        Identify the top 5 regulatory risk areas based on:
        1. Potential impact of non-compliance
        2. Complexity of implementation
        3. Areas where the company may have gaps based on the "May Apply" regulations

        For each risk area:
        1. Describe the risk in business terms
        2. Provide a risk rating (High, Medium, Low)
        3. Suggest risk mitigation strategies

        Here are the regulations that apply:
        {json.dumps([{'Regulation Number': r['Regulation Number'], 
                     'Regulation Title': r['Regulation Title'], 
                     'Reason for Applicability': r.get('Reason for Applicability', 'Not specified')} 
                    for r in applies[:10]], indent=2)}
        
        Here are regulations that may apply and require further review:
        {json.dumps([{'Regulation Number': r['Regulation Number'], 
                     'Regulation Title': r['Regulation Title'], 
                     'Reason for Applicability': r.get('Reason for Applicability', 'Not specified')} 
                    for r in may_apply[:10]], indent=2)}
        
        Your response should be structured, concise, and focused on strategic risk insights.
        """
        
        # Check cache first
        cache_key = self.get_cache_key(prompt)
        cached_response = self.check_cache(cache_key)
        if cached_response:
            print("Using cached risk assessment")
            return cached_response
        
        # Make API call to Claude 3.7 Sonnet
        print("Generating risk assessment using Claude 3.7 Sonnet...")
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "anthropic/claude-3-sonnet-20240229",
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 3000
        }
        
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=data
        )
        
        if response.status_code != 200:
            raise Exception(f"API request failed with status code {response.status_code}: {response.text}")
        
        result = response.json()
        risk_assessment = result['choices'][0]['message']['content']
        
        # Save to cache
        self.save_to_cache(cache_key, risk_assessment)
        
        return risk_assessment
    
    def generate_action_plan(self, analyzed_regulations):
        """
        Generate a strategic action plan based on the regulations using Claude 3.7 Sonnet.
        
        Args:
            analyzed_regulations (list): The analyzed regulations
            
        Returns:
            dict: The action plan
        """
        # Group regulations by applicability
        applies = [r for r in analyzed_regulations if r.get('Applicability') == 'Applies']
        may_apply = [r for r in analyzed_regulations if r.get('Applicability') == 'May Apply - Requires Further Review']
        
        # Prepare the prompt
        prompt = f"""
        You are a financial regulatory implementation expert. Based on the following analysis of regulations that apply to a financial institution, 
        provide a strategic action plan for regulatory compliance.

        Create a phased implementation approach with:
        1. Immediate actions (next 30 days)
        2. Short-term actions (1-3 months)
        3. Medium-term actions (3-6 months)
        4. Long-term actions (6-12 months)

        For each phase:
        1. Identify the key priorities
        2. Suggest specific implementation steps
        3. Identify key stakeholders who should be involved
        4. Suggest metrics to track progress

        Here are the regulations that apply:
        {json.dumps([{'Regulation Number': r['Regulation Number'], 
                     'Regulation Title': r['Regulation Title'], 
                     'Reason for Applicability': r.get('Reason for Applicability', 'Not specified')} 
                    for r in applies[:10]], indent=2)}
        
        Here are regulations that may apply and require further review:
        {json.dumps([{'Regulation Number': r['Regulation Number'], 
                     'Regulation Title': r['Regulation Title'], 
                     'Reason for Applicability': r.get('Reason for Applicability', 'Not specified')} 
                    for r in may_apply[:10]], indent=2)}
        
        Your response should be structured, concise, and focused on actionable implementation steps.
        """
        
        # Check cache first
        cache_key = self.get_cache_key(prompt)
        cached_response = self.check_cache(cache_key)
        if cached_response:
            print("Using cached action plan")
            return cached_response
        
        # Make API call to Claude 3.7 Sonnet
        print("Generating action plan using Claude 3.7 Sonnet...")
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "anthropic/claude-3-sonnet-20240229",
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 3000
        }
        
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=data
        )
        
        if response.status_code != 200:
            raise Exception(f"API request failed with status code {response.status_code}: {response.text}")
        
        result = response.json()
        action_plan = result['choices'][0]['message']['content']
        
        # Save to cache
        self.save_to_cache(cache_key, action_plan)
        
        return action_plan
    
    def generate_summary_document(self):
        """
        Generate a high-level summary document of the regulatory analysis.
        
        Returns:
            Path: Path to the saved Word document
        """
        # Load the analyzed regulations
        analyzed_regulations = self.load_analyzed_regulations()
        
        # Generate the thematic analysis, risk assessment, and action plan
        # Add a small delay between API calls to avoid rate limiting
        thematic_analysis = self.generate_thematic_analysis(analyzed_regulations)
        time.sleep(2)
        risk_assessment = self.generate_risk_assessment(analyzed_regulations)
        time.sleep(2)
        action_plan = self.generate_action_plan(analyzed_regulations)
        
        # Create a new Word document
        doc = docx.Document()
        
        # Set styles for professional formatting
        styles = doc.styles
        style = styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = docx.shared.Pt(11)
        
        # Add title page
        doc.add_heading('Regulatory Compliance', 0)
        doc.add_heading('Strategic Analysis & Action Plan', 1)
        
        # Add company name placeholder
        p = doc.add_paragraph()
        p.add_run('Prepared for: ').bold = True
        p.add_run('AIB Life')
        
        # Add date
        p = doc.add_paragraph()
        p.add_run('Date: ').bold = True
        p.add_run(f'{datetime.now().strftime("%B %d, %Y")}')
        
        # Add confidentiality notice
        p = doc.add_paragraph()
        p.add_run('CONFIDENTIAL').bold = True
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a page break
        doc.add_page_break()
        
        # Add table of contents header
        doc.add_heading('Table of Contents', 1)
        toc = doc.add_paragraph()
        toc.add_run('1. Executive Summary').bold = True
        toc.add_run('\n   • Key Findings\n   • Regulatory Landscape\n   • Compliance Posture\n\n')
        
        toc.add_run('2. Thematic Analysis').bold = True
        toc.add_run('\n   • Key Regulatory Themes\n   • Requirements by Theme\n   • Implementation Considerations\n\n')
        
        toc.add_run('3. Risk Assessment').bold = True
        toc.add_run('\n   • Top Regulatory Risks\n   • Risk Ratings\n   • Mitigation Strategies\n\n')
        
        toc.add_run('4. Strategic Action Plan').bold = True
        toc.add_run('\n   • Immediate Actions (30 days)\n   • Short-term Actions (1-3 months)\n   • Medium-term Actions (3-6 months)\n   • Long-term Actions (6-12 months)\n\n')
        
        toc.add_run('5. Appendices').bold = True
        toc.add_run('\n   • Regulatory Statistics\n   • Methodology\n   • Glossary')
        
        # Add a page break
        doc.add_page_break()
        
        # Add executive summary
        doc.add_heading('1. Executive Summary', 1)
        
        # Group regulations by applicability
        applies = [r for r in analyzed_regulations if r.get('Applicability') == 'Applies']
        does_not_apply = [r for r in analyzed_regulations if r.get('Applicability') == 'Does Not Apply']
        may_apply = [r for r in analyzed_regulations if r.get('Applicability') == 'May Apply - Requires Further Review']
        
        # Add key findings section with professional formatting
        doc.add_heading('Key Findings', 2)
        
        # Create a styled table for summary statistics
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Category'
        header_cells[1].text = 'Count (Percentage)'
        for cell in header_cells:
            # Make text bold
            cell.paragraphs[0].runs[0].bold = True
        
        # Add data rows
        data_rows = [
            ['Total Regulations', f"{len(analyzed_regulations)}"],
            ['Applicable Regulations', f"{len(applies)} ({len(applies)/len(analyzed_regulations)*100:.1f}%)"],
            ['May Apply - Requires Review', f"{len(may_apply)} ({len(may_apply)/len(analyzed_regulations)*100:.1f}%)"],
            ['Non-Applicable Regulations', f"{len(does_not_apply)} ({len(does_not_apply)/len(analyzed_regulations)*100:.1f}%)"],
        ]
        
        for i, (category, count) in enumerate(data_rows):
            row = table.rows[i]
            row.cells[0].text = category
            row.cells[1].text = count
        
        # Add a brief executive overview
        doc.add_heading('Regulatory Landscape', 2)
        doc.add_paragraph('This report provides a strategic analysis of the Central Bank regulations (Section 17A and Section 48) applicable to AIB Life based on a comprehensive review of the company\'s annual report and regulatory documentation.')
        
        # Add compliance posture summary
        doc.add_heading('Compliance Posture', 2)
        posture_para = doc.add_paragraph()
        posture_para.add_run('Overall Compliance Status: ').bold = True
        if len(applies) / len(analyzed_regulations) < 0.3:
            posture_para.add_run('Low Regulatory Exposure')
        elif len(applies) / len(analyzed_regulations) < 0.6:
            posture_para.add_run('Moderate Regulatory Exposure')
        else:
            posture_para.add_run('High Regulatory Exposure')
        
        # Add a page break
        doc.add_page_break()
        
        # Add thematic analysis with better formatting
        doc.add_heading('2. Thematic Analysis', 1)
        doc.add_heading('Key Regulatory Themes', 2)
        
        # Format the thematic analysis with better structure
        # Split the analysis into paragraphs and add proper formatting
        thematic_paragraphs = thematic_analysis.split('\n\n')
        for para in thematic_paragraphs:
            if para.strip().upper() == para.strip():
                # This is likely a header
                doc.add_heading(para.strip(), 3)
            elif ':' in para and para.split(':')[0].isupper():
                # This is likely a subheading with content
                parts = para.split(':', 1)
                p = doc.add_paragraph()
                p.add_run(parts[0] + ':').bold = True
                p.add_run(parts[1])
            else:
                # Regular paragraph
                doc.add_paragraph(para)
        
        # Add a page break
        doc.add_page_break()
        
        # Add risk assessment with better formatting
        doc.add_heading('3. Risk Assessment', 1)
        doc.add_heading('Top Regulatory Risks', 2)
        
        # Create a styled table for risk assessment
        risk_table = doc.add_table(rows=1, cols=3)
        risk_table.style = 'Table Grid'
        
        # Add header row
        risk_header = risk_table.rows[0].cells
        risk_header[0].text = 'Risk Area'
        risk_header[1].text = 'Risk Rating'
        risk_header[2].text = 'Key Mitigation Strategies'
        
        for cell in risk_header:
            cell.paragraphs[0].runs[0].bold = True
        
        # Parse risk assessment to extract structured data
        # This is a simplified approach - in reality, we'd need more sophisticated parsing
        risk_sections = risk_assessment.split('\n\n')
        risk_areas = []
        current_risk = {}
        
        for section in risk_sections:
            if section.strip().startswith('Risk Area'):
                if current_risk and 'name' in current_risk:
                    risk_areas.append(current_risk)
                current_risk = {'name': section.strip().replace('Risk Area ', '').split(':')[0]}
            elif 'Risk Rating' in section:
                current_risk['rating'] = section.split('Risk Rating:')[1].strip().split('\n')[0]
            elif 'Mitigation Strategies' in section:
                current_risk['mitigation'] = section.split('Mitigation Strategies:')[1].strip()
        
        # Add the last risk area if it exists
        if current_risk and 'name' in current_risk:
            risk_areas.append(current_risk)
        
        # If we couldn't parse the risks properly, just add the full text
        if not risk_areas:
            doc.add_paragraph(risk_assessment)
        else:
            # Add risk areas to the table
            for risk in risk_areas:
                row = risk_table.add_row()
                row.cells[0].text = risk.get('name', 'N/A')
                row.cells[1].text = risk.get('rating', 'N/A')
                row.cells[2].text = risk.get('mitigation', 'N/A')
        
        # Add the full risk assessment text after the table
        doc.add_heading('Detailed Risk Analysis', 2)
        doc.add_paragraph(risk_assessment)
        
        # Add a page break
        doc.add_page_break()
        
        # Add action plan with better formatting
        doc.add_heading('4. Strategic Action Plan', 1)
        
        # Create a timeline visualization
        doc.add_heading('Implementation Timeline', 2)
        timeline_table = doc.add_table(rows=5, cols=1)
        timeline_table.style = 'Table Grid'
        
        # Add timeline headers with colored backgrounds
        timeline_headers = [
            ('Immediate Actions (Next 30 Days)', 'FF9999'),  # Light red
            ('Short-term Actions (1-3 Months)', 'FFCC99'),   # Light orange
            ('Medium-term Actions (3-6 Months)', 'FFFFCC'),  # Light yellow
            ('Long-term Actions (6-12 Months)', 'CCFFCC'),   # Light green
            ('Ongoing Monitoring', 'CCCCFF')                 # Light blue
        ]
        
        for i, (header, _) in enumerate(timeline_headers):
            cell = timeline_table.rows[i].cells[0]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        # Add the detailed action plan
        doc.add_heading('Detailed Action Plan', 2)
        doc.add_paragraph(action_plan)
        
        # Add a page break
        doc.add_page_break()
        
        # Add appendices
        doc.add_heading('5. Appendices', 1)
        
        # Add statistical overview
        doc.add_heading('Regulatory Statistics', 2)
        
        # Add section type statistics with applicability breakdown
        section_17a = [r for r in analyzed_regulations if r.get('Section Type') == '17A']
        section_48 = [r for r in analyzed_regulations if r.get('Section Type') == '48']
        
        # Count applicability for each section type
        section_17a_applies = [r for r in section_17a if r.get('Applicability') == 'Applies']
        section_17a_may_apply = [r for r in section_17a if r.get('Applicability') == 'May Apply - Requires Further Review']
        section_17a_not_apply = [r for r in section_17a if r.get('Applicability') == 'Does Not Apply']
        
        section_48_applies = [r for r in section_48 if r.get('Applicability') == 'Applies']
        section_48_may_apply = [r for r in section_48 if r.get('Applicability') == 'May Apply - Requires Further Review']
        section_48_not_apply = [r for r in section_48 if r.get('Applicability') == 'Does Not Apply']
        
        # Create a table for section statistics with applicability
        section_table = doc.add_table(rows=3, cols=5)
        section_table.style = 'Table Grid'
        
        # Add header
        section_header = section_table.rows[0].cells
        section_header[0].text = 'Regulation Type'
        section_header[1].text = 'Total Count'
        section_header[2].text = 'Applies'
        section_header[3].text = 'May Apply'
        section_header[4].text = 'Does Not Apply'
        
        for cell in section_header:
            cell.paragraphs[0].runs[0].bold = True
        
        # Add data for Section 17A
        section_table.rows[1].cells[0].text = 'Section 17A'
        section_table.rows[1].cells[1].text = str(len(section_17a))
        section_table.rows[1].cells[2].text = f"{len(section_17a_applies)} ({len(section_17a_applies)/len(section_17a)*100:.1f}%)"
        section_table.rows[1].cells[3].text = f"{len(section_17a_may_apply)} ({len(section_17a_may_apply)/len(section_17a)*100:.1f}%)"
        section_table.rows[1].cells[4].text = f"{len(section_17a_not_apply)} ({len(section_17a_not_apply)/len(section_17a)*100:.1f}%)"
        
        # Add data for Section 48
        section_table.rows[2].cells[0].text = 'Section 48'
        section_table.rows[2].cells[1].text = str(len(section_48))
        section_table.rows[2].cells[2].text = f"{len(section_48_applies)} ({len(section_48_applies)/len(section_48)*100:.1f}%)"
        section_table.rows[2].cells[3].text = f"{len(section_48_may_apply)} ({len(section_48_may_apply)/len(section_48)*100:.1f}%)"
        section_table.rows[2].cells[4].text = f"{len(section_48_not_apply)} ({len(section_48_not_apply)/len(section_48)*100:.1f}%)"
        
        # Add part statistics
        doc.add_heading('Regulations by Part', 3)
        
        parts = {}
        for reg in analyzed_regulations:
            key = f"{reg['Section Type']} - Part {reg['Part Number']}: {reg['Part Name']}"
            if key not in parts:
                parts[key] = []
            parts[key].append(reg)
        
        # Create a table for part statistics with applicability
        part_table = doc.add_table(rows=1, cols=5)
        part_table.style = 'Table Grid'
        
        # Add header
        part_header = part_table.rows[0].cells
        part_header[0].text = 'Regulation Part'
        part_header[1].text = 'Total Count'
        part_header[2].text = 'Applies'
        part_header[3].text = 'May Apply'
        part_header[4].text = 'Does Not Apply'
        
        for cell in part_header:
            cell.paragraphs[0].runs[0].bold = True
        
        # Add data rows with applicability breakdown
        for part_key, part_regs in parts.items():
            row = part_table.add_row()
            row.cells[0].text = part_key
            row.cells[1].text = str(len(part_regs))
            
            # Count applicability for this part
            applies = [r for r in part_regs if r.get('Applicability') == 'Applies']
            may_apply = [r for r in part_regs if r.get('Applicability') == 'May Apply - Requires Further Review']
            not_apply = [r for r in part_regs if r.get('Applicability') == 'Does Not Apply']
            
            # Add applicability counts and percentages
            row.cells[2].text = f"{len(applies)} ({len(applies)/len(part_regs)*100:.1f}%)"
            row.cells[3].text = f"{len(may_apply)} ({len(may_apply)/len(part_regs)*100:.1f}%)"
            row.cells[4].text = f"{len(not_apply)} ({len(not_apply)/len(part_regs)*100:.1f}%)"
        
        # Add methodology section
        doc.add_heading('Methodology', 2)
        method_para = doc.add_paragraph()
        method_para.add_run('Analysis Approach: ').bold = True
        method_para.add_run('This analysis was conducted using Claude 3.7 Sonnet, an advanced AI system, to evaluate the applicability of Central Bank regulations to AIB Life based on the company\'s annual report.')
        
        method_para = doc.add_paragraph()
        method_para.add_run('Data Sources: ').bold = True
        method_para.add_run('The analysis is based on the following sources:')
        
        sources = doc.add_paragraph(style='List Bullet')
        sources.add_run('Central Bank Reform Act 2010 (Section 17A) regulations')
        sources = doc.add_paragraph(style='List Bullet')
        sources.add_run('Central Bank (Supervision and Enforcement) Act 2013 (Section 48) regulations')
        sources = doc.add_paragraph(style='List Bullet')
        sources.add_run('AIB Life Annual Report')
        
        # Add glossary
        doc.add_heading('Glossary', 2)
        glossary_terms = [
            ('Applies', 'Regulation directly applies to the company based on its activities and structure.'),
            ('May Apply', 'Regulation potentially applies but requires further review to determine applicability.'),
            ('Does Not Apply', 'Regulation does not apply to the company based on its activities and structure.'),
            ('Section 17A', 'Regulations under Central Bank Reform Act 2010 (Section 17A).'),
            ('Section 48', 'Regulations under Central Bank (Supervision and Enforcement) Act 2013 (Section 48).')
        ]
        
        glossary_table = doc.add_table(rows=1, cols=2)
        glossary_table.style = 'Table Grid'
        
        # Add header
        glossary_header = glossary_table.rows[0].cells
        glossary_header[0].text = 'Term'
        glossary_header[1].text = 'Definition'
        
        for cell in glossary_header:
            cell.paragraphs[0].runs[0].bold = True
        
        # Add glossary terms
        for term, definition in glossary_terms:
            row = glossary_table.add_row()
            row.cells[0].text = term
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = definition
        
        # Save the document with a new filename to avoid permission issues
        output_path = self.output_dir / "regulatory_strategic_summary.docx"
        doc.save(output_path)
        
        print(f"Strategic summary document saved to {output_path}")
        return output_path

def main():
    """Run the high-level summary generator."""
    print("Note: Using the provided OpenRouter API key for Claude 3.7 Sonnet")
    generator = HighLevelSummaryGenerator()
    generator.generate_summary_document()

if __name__ == "__main__":
    main()
